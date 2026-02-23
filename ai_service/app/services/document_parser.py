"""
文档解析服务
支持PDF、DOCX、TXT、MD等格式
"""

import hashlib
import os
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional, BinaryIO

from app.core.logging import get_logger
from app.core.config import settings

logger = get_logger(__name__)


@dataclass
class ParsedDocument:
    """解析后的文档"""
    file_path: str
    file_name: str
    file_hash: str
    file_size: int
    mime_type: str
    content: str
    metadata: dict
    extracted_at: datetime


class BaseParser(ABC):
    """文档解析器基类"""
    
    supported_extensions: List[str] = []
    
    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """解析文档"""
        pass
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否支持该文件"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions


class PDFParser(BaseParser):
    """PDF解析器"""
    
    supported_extensions = ['.pdf']
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析PDF文件"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            
            # 提取文本
            text_parts = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"\n--- Page {i + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning("pdf.page_extract_failed", page=i, error=str(e))
            
            content = "\n".join(text_parts)
            
            # 获取元数据
            pdf_metadata = reader.metadata or {}
            
            # 计算文件hash
            file_hash = self._calculate_hash(file_path)
            file_size = os.path.getsize(file_path)
            
            return ParsedDocument(
                file_path=file_path,
                file_name=Path(file_path).name,
                file_hash=file_hash,
                file_size=file_size,
                mime_type="application/pdf",
                content=content,
                metadata={
                    "page_count": len(reader.pages),
                    "title": pdf_metadata.get("/Title", ""),
                    "author": pdf_metadata.get("/Author", ""),
                    "subject": pdf_metadata.get("/Subject", ""),
                },
                extracted_at=datetime.now(),
            )
            
        except Exception as e:
            logger.error("pdf.parse_failed", file_path=file_path, error=str(e))
            raise
    
    def _calculate_hash(self, file_path: str) -> str:
        """计算文件SHA256哈希"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()


class DOCXParser(BaseParser):
    """DOCX解析器"""
    
    supported_extensions = ['.docx']
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析DOCX文件"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格内容
            tables_content = []
            for table_idx, table in enumerate(doc.tables):
                table_rows = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(row_text))
                if table_rows:
                    tables_content.append(f"\n--- Table {table_idx + 1} ---\n" + "\n".join(table_rows))
            
            content = "\n\n".join(paragraphs)
            if tables_content:
                content += "\n\n" + "\n\n".join(tables_content)
            
            # 获取文档属性
            core_props = doc.core_properties
            
            file_hash = self._calculate_hash(file_path)
            file_size = os.path.getsize(file_path)
            
            return ParsedDocument(
                file_path=file_path,
                file_name=Path(file_path).name,
                file_hash=file_hash,
                file_size=file_size,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content=content,
                metadata={
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                },
                extracted_at=datetime.now(),
            )
            
        except Exception as e:
            logger.error("docx.parse_failed", file_path=file_path, error=str(e))
            raise
    
    def _calculate_hash(self, file_path: str) -> str:
        """计算文件SHA256哈希"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()


class TXTParser(BaseParser):
    """文本文件解析器（支持txt/md）"""
    
    supported_extensions = ['.txt', '.md', '.markdown']
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析文本文件"""
        try:
            # 检测编码
            encoding = self._detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read()
            
            # 检测是否为Markdown
            ext = Path(file_path).suffix.lower()
            is_markdown = ext in ['.md', '.markdown']
            
            # Markdown元数据提取（Front Matter）
            metadata = {}
            if is_markdown:
                metadata = self._extract_markdown_metadata(content)
                content = self._strip_front_matter(content)
            
            file_hash = self._calculate_hash(file_path)
            file_size = os.path.getsize(file_path)
            
            mime_type = "text/markdown" if is_markdown else "text/plain"
            
            return ParsedDocument(
                file_path=file_path,
                file_name=Path(file_path).name,
                file_hash=file_hash,
                file_size=file_size,
                mime_type=mime_type,
                content=content,
                metadata={
                    "encoding": encoding,
                    "is_markdown": is_markdown,
                    **metadata,
                },
                extracted_at=datetime.now(),
            )
            
        except Exception as e:
            logger.error("txt.parse_failed", file_path=file_path, error=str(e))
            raise
    
    def _detect_encoding(self, file_path: str) -> str:
        """检测文件编码"""
        import chardet
        
        with open(file_path, 'rb') as f:
            raw_data = f.read(10000)  # 读取前10KB检测
            result = chardet.detect(raw_data)
            encoding = result.get('encoding', 'utf-8')
            # 处理GB2312 -> GB18030
            if encoding and encoding.lower() == 'gb2312':
                encoding = 'gb18030'
            return encoding or 'utf-8'
    
    def _extract_markdown_metadata(self, content: str) -> dict:
        """提取Markdown front matter"""
        if content.startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                try:
                    import yaml
                    metadata = yaml.safe_load(parts[1])
                    return metadata or {}
                except Exception:
                    pass
        return {}
    
    def _strip_front_matter(self, content: str) -> str:
        """去除front matter"""
        if content.startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                return parts[2].strip()
        return content
    
    def _calculate_hash(self, file_path: str) -> str:
        """计算文件SHA256哈希"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()


class DocumentParserService:
    """文档解析服务"""
    
    def __init__(self):
        self.parsers: List[BaseParser] = [
            PDFParser(),
            DOCXParser(),
            TXTParser(),
        ]
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文档
        
        Args:
            file_path: 文件路径
        
        Returns:
            ParsedDocument: 解析后的文档
        """
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 找到合适的解析器
        for parser in self.parsers:
            if parser.can_parse(file_path):
                logger.info("parser.selected", file_path=file_path, parser=parser.__class__.__name__)
                return parser.parse(file_path)
        
        # 没有合适的解析器
        raise ValueError(f"不支持的文件格式: {file_path}")
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否支持解析该文件"""
        return any(parser.can_parse(file_path) for parser in self.parsers)
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名列表"""
        extensions = []
        for parser in self.parsers:
            extensions.extend(parser.supported_extensions)
        return extensions


# 全局解析服务实例
_parser_service: Optional[DocumentParserService] = None


def get_parser_service() -> DocumentParserService:
    """获取文档解析服务单例"""
    global _parser_service
    if _parser_service is None:
        _parser_service = DocumentParserService()
    return _parser_service
